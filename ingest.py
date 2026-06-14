import os
import fitz
import docx
from dotenv import load_dotenv
from tqdm import tqdm

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma

from config import settings

load_dotenv()


def read_pdf(file_path: str) -> tuple[str, dict]:
    pages = []
    with fitz.open(file_path) as pdf:
        page_count = len(pdf)
        for page_num, page in enumerate(pdf, start=1):
            pages.append(page.get_text())
    return "\n\n".join(pages), {"page_count": page_count}


def read_docx(file_path: str) -> tuple[str, dict]:
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs), {}


def read_txt(file_path: str) -> tuple[str, dict]:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read(), {}


READERS = {
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".txt": read_txt,
}


def load_documents() -> list[Document]:
    documents = []
    files = [
        f for f in os.listdir(settings.documents_dir)
        if os.path.splitext(f)[1].lower() in READERS
    ]

    if not files:
        return documents

    for filename in tqdm(files, desc="Loading documents"):
        ext = os.path.splitext(filename)[1].lower()
        file_path = os.path.join(settings.documents_dir, filename)
        text, extra_meta = READERS[ext](file_path)
        documents.append(
            Document(
                page_content=text,
                metadata={"source": filename, "path": file_path, **extra_meta},
            )
        )

    return documents


def ingest(documents_dir: str | None = None) -> int:
    if documents_dir:
        settings.documents_dir = documents_dir

    print(f"Loading documents from '{settings.documents_dir}'...")
    documents = load_documents()

    if not documents:
        print("No supported documents found (.pdf, .docx, .txt).")
        return 0

    print(f"Splitting {len(documents)} document(s) into chunks...")
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
    )
    chunks = splitter.split_documents(documents)
    print(f"  Created {len(chunks)} chunks.")

    print("Embedding and storing chunks...")
    embeddings = HuggingFaceEmbeddings(model_name=settings.embedding_model)

    with tqdm(total=len(chunks), desc="Embedding") as pbar:
        Chroma.from_documents(
            documents=chunks,
            embedding=embeddings,
            persist_directory=settings.vector_db_dir,
        )
        pbar.update(len(chunks))

    print(f"Done. Vector DB saved to '{settings.vector_db_dir}'.")
    return len(chunks)


if __name__ == "__main__":
    ingest()
